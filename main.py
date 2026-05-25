"""
TailoredResume.ai - Agentic AI Resume Tailoring Backend
======================================================
Author: Nikhil Pitroda
Version: 1.1.0

This module provides the core logic for an agentic AI-driven resume tailoring platform.
It leverages Google Gemini AI models to analyze job descriptions and optimize candidate
resumes for ATS compliance and professional excellence.

Core Features:
- JD Analysis: Extracting keywords, priorities, and requirements.
- Agentic Tailoring: Iterative resume generation with feedback loops.
- ATS Scoring: Evaluating resumes based on keyword match, quantification, and diversity.
- PDF Generation: Creating professional, ATS-optimized PDF documents.
"""

import os
import json
import re
import time
import google.generativeai as genai
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
from datetime import datetime
from dotenv import load_dotenv

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle, Indenter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import Color

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

app = FastAPI(
    title="TailoredResume.ai API",
    description="Agentic AI Resume Tailoring Service using Google Gemini",
    version="1.1.0"
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
STATIC_DIR = os.path.join(BASE_DIR, "static")

api_key = os.getenv("GEMINI_API_KEY")
if api_key:
    genai.configure(api_key=api_key)

MODELS_TO_TRY = [
        'gemini-3-flash-preview',
        'gemini-2.5-flash-lite',
        'gemini-3.1-flash-lite-preview',
        'gemini-2.0-flash'
]

import io

def create_pdf(data: dict, output_buffer, template="classic"):
    """Generates an ATS-optimized PDF resume from structured JSON data.

    This function uses the ReportLab library to build a professional resume layout.
    It handles personal information, professional summaries, technical skills,
    work experience, personal projects, education, and certifications.

    Design Goals:
    - Maximum Readability: Using standard fonts and clear hierarchies.
    - ATS Compliance: Avoiding complex graphics that confuse parsers.
    - Professional Aesthetic: Clean lines, consistent spacing, and bold highlights.

    Args:
        data (dict): The structured resume data containing personal information,
            summary, skills, experience, projects, and education. Expected keys:
            - "personal_info": dict with name, email, phone, linkedin, etc.
            - "summary": string of professional summary.
            - "skills": dict or list of technical skills.
            - "experience": list of work history entries.
            - "projects": list of project entries.
            - "education": list of academic history.
            - "certifications": list of professional certifications.
        output_buffer (io.BytesIO): The in-memory buffer where the PDF content 
            will be written for streaming or saving.

    Returns:
        None: The PDF is written directly to the output_buffer.
    """
    doc = SimpleDocTemplate(
        output_buffer,
        pagesize=letter,
        rightMargin=22,
        leftMargin=22,
        topMargin=12,
        bottomMargin=5
    )

    styles = getSampleStyleSheet()
    
    # Determine styles based on template
    font_name = "Helvetica"
    font_bold = "Helvetica-Bold"
    font_italic = "Helvetica-Oblique"
    line_color = Color(0, 0, 0)
    
    if template == "modern":
        line_color = Color(0.2, 0.4, 0.8)
    elif template == "creative":
        line_color = Color(0.5, 0.2, 0.7)
    elif template == "executive":
        font_name = "Times-Roman"
        font_bold = "Times-Bold"
        font_italic = "Times-Italic"
        line_color = Color(0.15, 0.25, 0.45)
    elif template == "split":
        line_color = Color(0.2, 0.4, 0.5)

    styles.add(ParagraphStyle(
        name='CompName', 
        parent=styles['Heading1'], 
        fontName=font_bold,
        fontSize=24 if template == "creative" else (21 if template == "executive" else (22 if template == "modern" else 18)), 
        spaceAfter=2, 
        alignment=0 if template in ["modern", "split"] else 1, 
        textColor=line_color
    ))
    
    styles.add(ParagraphStyle(
        name='CompSectionHeader', 
        parent=styles['Heading2'], 
        fontName=font_bold,
        fontSize=11 if template in ["modern", "creative", "split"] else 10.5, 
        spaceAfter=0.2, 
        spaceBefore=3, 
        textColor=line_color
    ))
    
    styles.add(ParagraphStyle(
        name='CompContact', 
        parent=styles['Normal'], 
        fontName=font_name,
        fontSize=8.5, 
        spaceAfter=3, 
        alignment=0 if template in ["modern", "split"] else 1, 
        textColor=Color(0.15, 0.15, 0.15)
    ))
    
    styles.add(ParagraphStyle(name='CompSummary', parent=styles['Normal'], fontName=font_name, fontSize=9, spaceAfter=0.5, leading=10))
    styles.add(ParagraphStyle(name='CompItemTitle', parent=styles['Normal'], fontName=font_bold, fontSize=9.5, spaceAfter=0.2, spaceBefore=0.5))
    styles.add(ParagraphStyle(name='CompItemDate', parent=styles['Normal'], fontName=font_italic, fontSize=8.5, alignment=2, textColor=Color(0.2, 0.2, 0.2)))
    styles.add(ParagraphStyle(name='CompResumeBullet', parent=styles['Normal'], fontName=font_name, fontSize=8.6, leading=9.6, leftIndent=10, bulletIndent=3, spaceAfter=0.3))
    styles.add(ParagraphStyle(name='CompContent', parent=styles['Normal'], fontName=font_name, fontSize=8.6, leading=9.6))

    elements = []
    line = HRFlowable(width="100%", thickness=1.0, color=line_color, spaceAfter=1.5, spaceBefore=0.5)

    def create_header_table(left_text, right_text, left_w=None, right_w=None):
        col_widths = ['78%', '22%'] if (left_w is None or right_w is None) else [left_w, right_w]
        t = Table(
            [[Paragraph(left_text, styles['CompItemTitle']), Paragraph(right_text, styles['CompItemDate'])]], 
            colWidths=col_widths
        )
        t.setStyle(TableStyle([
            ('LEFTPADDING', (0,0), (-1,-1), 0),
            ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ('TOPPADDING', (0,0), (-1,-1), 0),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0),
        ]))
        return t

    # 1. Personal Info
    pi = data.get("personal_info", {})
    name = pi.get("name", "RESUME").upper()
    elements.append(Paragraph(f"<b>{name}</b>", styles['CompName']))
    
    contacts = []
    if pi.get("location"):
        contacts.append(pi.get("location"))

    for field in ["phone", "email", "linkedin", "github", "portfolio", "kaggle"]:
        val = pi.get(field)
        if val:
            if "@" in val: # email
                contacts.append(f'<a href="mailto:{val}">{val}</a>')
            elif "." in val: # url
                link = val if val.startswith("http") else f"https://{val}"
                display = val.replace("https://", "").replace("www.", "")
                contacts.append(f'<a href="{link}">{display}</a>')
            else:
                contacts.append(val)
    
    if template == "executive":
        # Add a top line above contact details
        elements.append(HRFlowable(width="100%", thickness=0.8, color=line_color, spaceAfter=2.5, spaceBefore=2))
        elements.append(Paragraph(" | ".join(contacts), styles['CompContact']))
        # Add a bottom line below contact details
        elements.append(HRFlowable(width="100%", thickness=0.8, color=line_color, spaceAfter=6, spaceBefore=1.5))
    else:
        elements.append(Paragraph(" | ".join(contacts), styles['CompContact']))

    if template == "split":
        # Two-column layout
        left_elements = []
        right_elements = []
        
        # Build Left Side (Sidebar)
        # 1. Contact info listed vertically
        left_elements.append(Paragraph("<b>CONTACT</b>", styles['CompSectionHeader']))
        left_elements.append(HRFlowable(width="100%", thickness=0.8, color=line_color, spaceAfter=3, spaceBefore=1))
        
        if pi.get("location"):
            left_elements.append(Paragraph(f"<b>Location:</b> {pi.get('location')}", styles['CompContact']))
        if pi.get("phone"):
            left_elements.append(Paragraph(f"<b>Phone:</b> {pi.get('phone')}", styles['CompContact']))
        if pi.get("email"):
            left_elements.append(Paragraph(f'<b>Email:</b> <a href="mailto:{pi.get("email")}">{pi.get("email")}</a>', styles['CompContact']))
        
        for field in ["linkedin", "github", "portfolio", "kaggle"]:
            val = pi.get(field)
            if val:
                link = val if val.startswith("http") else f"https://{val}"
                display = val.replace("https://", "").replace("www.", "")
                if len(display) > 26:
                    display = display[:23] + "..."
                left_elements.append(Paragraph(f'<b>{field.capitalize()}:</b> <a href="{link}">{display}</a>', styles['CompContact']))
                
        left_elements.append(Spacer(1, 8))
        
        # 2. Education
        if data.get("education"):
            left_elements.append(Paragraph("<b>EDUCATION</b>", styles['CompSectionHeader']))
            left_elements.append(HRFlowable(width="100%", thickness=0.8, color=line_color, spaceAfter=3, spaceBefore=1))
            for edu in data["education"]:
                inst = edu.get('institution', '')
                deg = edu.get('degree', '')
                dates = f"{edu.get('start_date', '')} - {edu.get('end_date', '')}"
                gpa = edu.get('cgpa') or edu.get('gpa')
                
                left_elements.append(Paragraph(f"<b>{inst}</b>", styles['CompSummary']))
                left_elements.append(Paragraph(deg, styles['CompContact']))
                if gpa:
                    left_elements.append(Paragraph(f"GPA: {gpa}", styles['CompContact']))
                left_elements.append(Paragraph(f"<i>{dates}</i>", styles['CompContact']))
                left_elements.append(Spacer(1, 4))
                
        # 3. Skills
        if data.get("skills"):
            left_elements.append(Paragraph("<b>TECHNICAL SKILLS</b>", styles['CompSectionHeader']))
            left_elements.append(HRFlowable(width="100%", thickness=0.8, color=line_color, spaceAfter=3, spaceBefore=1))
            if isinstance(data["skills"], dict):
                for cat, items in data["skills"].items():
                    label = cat.replace('_', ' ')
                    left_elements.append(Paragraph(f"<b>{label}:</b>", styles['CompSummary']))
                    left_elements.append(Paragraph(", ".join(items), styles['CompContact']))
                    left_elements.append(Spacer(1, 3))
            left_elements.append(Spacer(1, 4))
            
        # 4. Certifications
        if data.get("certifications"):
            left_elements.append(Paragraph("<b>CERTIFICATIONS</b>", styles['CompSectionHeader']))
            left_elements.append(HRFlowable(width="100%", thickness=0.8, color=line_color, spaceAfter=3, spaceBefore=1))
            for cert in data["certifications"]:
                left_elements.append(Paragraph(f"• {cert}", styles['CompContact']))
                left_elements.append(Spacer(1, 2))
                
        # Build Right Side (Main Content)
        # 1. Summary
        if data.get("summary"):
            right_elements.append(Paragraph("<b>PROFESSIONAL SUMMARY</b>", styles['CompSectionHeader']))
            right_elements.append(HRFlowable(width="100%", thickness=0.8, color=line_color, spaceAfter=3, spaceBefore=1))
            right_elements.append(Paragraph(data["summary"], styles['CompSummary']))
            right_elements.append(Spacer(1, 8))
            
        # 2. Experience
        if data.get("experience"):
            right_elements.append(Paragraph("<b>EXPERIENCE</b>", styles['CompSectionHeader']))
            right_elements.append(HRFlowable(width="100%", thickness=0.8, color=line_color, spaceAfter=3, spaceBefore=1))
            for exp in data["experience"]:
                left_txt = f"<b>{exp.get('company', '')}</b> | {exp.get('role', '')}"
                right_txt = f"<i>{exp.get('start_date', '')} - {exp.get('end_date', '')}</i>"
                right_elements.append(create_header_table(left_txt, right_txt, 275, 95))
                for b in exp.get("bullet_points", []):
                    right_elements.append(Paragraph(f"• {b}", styles['CompResumeBullet']))
                right_elements.append(Spacer(1, 2))
                
        # 3. Projects
        if data.get("projects"):
            right_elements.append(Paragraph("<b>PROJECTS</b>", styles['CompSectionHeader']))
            right_elements.append(HRFlowable(width="100%", thickness=0.8, color=line_color, spaceAfter=3, spaceBefore=1))
            for p in data["projects"]:
                pname = p.get("name", "")
                techs = ", ".join(p.get("technologies", []))
                github = p.get("github_link", "")
                right_txt = ""
                if github:
                    link = github if github.startswith("http") else f"https://{github}"
                    right_txt = f'<a href="{link}"><i>GitHub</i></a>'
                right_elements.append(create_header_table(f"<b>{pname}</b> | <i>{techs}</i>", right_txt, 275, 95))
                for b in p.get("bullet_points", []):
                    right_elements.append(Paragraph(f"• {b}", styles['CompResumeBullet']))
                right_elements.append(Spacer(1, 2))
                
        # Assemble into Table
        grid_table = Table([[left_elements, Spacer(1, 1), right_elements]], colWidths=[175, 15, 378])
        grid_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
            ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ('TOPPADDING', (0,0), (-1,-1), 0),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0),
        ]))
        elements.append(grid_table)
    else:
        # Single column layout (classic, modern, creative, executive)
        sections_keys = ["summary", "skills", "experience", "projects", "education", "certifications"]
        sections_titles = ["PROFESSIONAL SUMMARY", "TECHNICAL SKILLS", "EXPERIENCE", "PROJECTS", "EDUCATION", "CERTIFICATIONS"]

        for i, key in enumerate(sections_keys):
            if not data.get(key): continue
            
            elements.append(Paragraph(f"<b>{sections_titles[i]}</b>", styles['CompSectionHeader']))
            elements.append(line)
            
            if key == "summary":
                elements.append(Paragraph(data[key], styles['CompSummary']))
            elif key == "skills":
                skills_data = []
                if isinstance(data[key], dict):
                    for cat, items in data[key].items():
                        label = cat.replace('_', ' ')
                        skills_data.append([
                            Paragraph(f"<b>{label}:</b>", styles['CompSummary']),
                            Paragraph(", ".join(items), styles['CompSummary'])
                        ])
                
                # Create a table for the skills section to align domains on the left
                # Total width is 612 (letter) - 22*2 (margins) = 568
                # Domain column: 135, Skills column: 433
                s_table = Table(skills_data, colWidths=[135, 433])
                s_table.setStyle(TableStyle([
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('LEFTPADDING', (0,0), (-1,-1), 0),
                    ('RIGHTPADDING', (0,0), (-1,-1), 0),
                    ('TOPPADDING', (0,0), (-1,-1), 0.5),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 0.5),
                ]))
                elements.append(Indenter(left=10))
                elements.append(s_table)
                elements.append(Indenter(left=-10))
                elements.append(Spacer(1, 1))
            elif key == "experience":
                for exp in data[key]:
                    left = f"<b>{exp.get('company', '')}</b> | {exp.get('role', '')}"
                    right = f"<i>{exp.get('start_date', '')} - {exp.get('end_date', '')}</i>"
                    elements.append(create_header_table(left, right))
                    for b in exp.get("bullet_points", []):
                        elements.append(Paragraph(f"• {b}", styles['CompResumeBullet']))
            elif key == "projects":
                for p in data[key]:
                    pname = p.get("name", "")
                    techs = ", ".join(p.get("technologies", []))
                    github = p.get("github_link", "").strip()
                    right_text = ""
                    if github:
                        link = github if github.startswith("http") else f"https://{github}"
                        right_text = f'<a href="{link}"><i>GitHub</i></a>'
                    elements.append(create_header_table(f"<b>{pname}</b> | <i>{techs}</i>", right_text))
                    for b in p.get("bullet_points", []):
                        elements.append(Paragraph(f"• {b}", styles['CompResumeBullet']))
            elif key == "education":
                for edu in data[key]:
                    left = f"<b>{edu.get('institution', '')}</b> | {edu.get('location', '')}"
                    right = f"<i>{edu.get('start_date', '')} - {edu.get('end_date', '')}</i>"
                    elements.append(create_header_table(left, right))
                    deg = f"{edu.get('degree', '')}"
                    gpa = edu.get('cgpa') or edu.get('gpa')
                    if gpa: deg += f" | GPA: {gpa}"
                    elements.append(Paragraph(deg, styles['CompContent']))
                    elements.append(Spacer(1, 1))
            elif key == "certifications":
                elements.append(Paragraph(", ".join(data[key]), styles['CompSummary']))

    def add_meta(canvas, doc):
        canvas.setTitle(f"Tailored Resume - TailoredResume.ai")
        canvas.setAuthor("TailoredResume.ai")
        canvas.setSubject("ATS Optimized Resume")

    doc.build(elements, onFirstPage=add_meta, onLaterPages=add_meta)

def check_quantification(bullet_points: List[str]) -> float:
    """Calculates the ratio of quantified bullet points in a resume section.

    A bullet point is considered "quantified" if it contains at least one digit (0-9).
    This is a critical metric for modern resumes as it demonstrates impact and results.

    Args:
        bullet_points (List[str]): A list of string representations of resume bullets.

    Returns:
        float: A value between 0.0 and 1.0. (e.g., 0.85 means 85% of bullets are quantified).
    """
    scored = 0
    for bullet in bullet_points:
        if re.search(r'\d', bullet): scored += 1
    return scored / len(bullet_points) if bullet_points else 0

def get_action_verbs(bullet_points: List[str]) -> List[str]:
    """Extracts the primary action verb from the beginning of each bullet point.

    Used to detect linguistic diversity and ensure that the candidate is using 
    strong, varied terminology rather than repeating "Developed" or "Worked on".

    Algorithm:
    1. Removes leading bullet symbols (*, -, •, digits).
    2. Takes the first word of the remaining string.
    3. Normalizes to lowercase and clears punctuation.

    Args:
        bullet_points (List[str]): A list of resume bullet points.

    Returns:
        List[str]: A list of lowercase verbs (one per bullet point).
    """
    verbs = []
    for bullet in bullet_points:
        clean = re.sub(r'^[•\-\*\d\.\s]+', '', bullet).strip()
        if clean: verbs.append(clean.split()[0].rstrip(',').lower())
    return verbs

def score_resume_internal(data: Dict[str, Any], analysis: Dict[str, Any]) -> Dict[str, Any]:
    """Evaluates the engineered resume against specific JD requirements using an ATS algorithm.

    This function simulates a high-tier Applicant Tracking System (ATS) by scoring the 
    resume on five key dimensions: Keyword Saturation, Impact Quantification, 
    Linguistic Diversity, Word Count Precision, and Skill Categorization.

    Scoring Weights:
    - Keywords: 30% (Critical for initial filters)
    - Quantification: 30% (Measures impact/results)
    - Diversity: 15% (Measures professional communication quality)
    - Word Count: 15% (Enforces conciseness and readability)
    - Skill Category: 10% (Ensures structured technical layout)

    Args:
        data (Dict[str, Any]): The generated resume JSON object.
        analysis (Dict[str, Any]): The structured JD analysis containing required skills.

    Returns:
        Dict[str, Any]: A report object with the final score, feedback messages, 
            and breakdown of individual category metrics.
    """
    all_text = json.dumps(data).lower()
    feedback = []
    scores = {}
    
    # 1. Keywords
    must_haves = analysis.get('must_have_skills', [])
    keywords_found = [k for k in must_haves if k.lower() in all_text]
    scores['keyword'] = len(keywords_found) / len(must_haves) if must_haves else 1.0

    # 2. Quantification
    all_bullets = []
    for exp in data.get('experience', []): all_bullets.extend(exp.get('bullet_points', []))
    for proj in data.get('projects', []): all_bullets.extend(proj.get('bullet_points', []))
    scores['quant'] = check_quantification(all_bullets)

    # 3. Diversity
    verbs = get_action_verbs(all_bullets)
    unique_verbs = set(verbs)
    scores['diversity'] = len(unique_verbs) / len(verbs) if verbs else 1.0
    if len(unique_verbs) < len(verbs):
        repeats = [v for v in unique_verbs if verbs.count(v) > 1]
        feedback.append(f"REPETITION DETECTED: Action verbs used more than once: {', '.join(repeats[:3])}")

    # 4. Word Count & Line Length (Strict 18-22 words for single-line bullets)
    wc_scores = []
    for b in all_bullets:
        wc = len(b.split())
        if 18 <= wc <= 22: wc_scores.append(1.0)
        elif 16 <= wc <= 24: wc_scores.append(0.5)
        else: wc_scores.append(0.0)
    scores['word_count'] = sum(wc_scores) / len(wc_scores) if wc_scores else 1.0

    # 5. Skills Cat & Bullet Count Consistency
    categories = list(data.get('skills', {}).keys())
    scores['skills_cat'] = min(len(categories) / 4, 1.0)
    
    # Check if exactly 4 bullets per project
    proj_bullet_counts = [len(p.get('bullet_points', [])) for p in data.get('projects', [])]
    if any(c != 4 for c in proj_bullet_counts):
        feedback.append("INCONSISTENT LAYOUT: Every project MUST have exactly 4 bullet points.")
        scores['word_count'] *= 0.8 # Penalty for layout inconsistency

    if scores['word_count'] < 0.9:
        feedback.append("CRITICAL: Bullet length variance. Ensure each bullet is strictly between 18-22 words to stay on one line.")
    if scores['quant'] < 0.9:
        feedback.append("IMPACT RATIO LOW: Every bullet must contain a quantifiable metric (%, $, #).")
    if len(keywords_found) < len(must_haves) * 0.9:
        missing = [k for k in must_haves if k.lower() not in all_text]
        feedback.append(f"MISSING CRITICAL TERMS: {', '.join(missing[:5])}")

    final_score = (scores['keyword'] * 30 + scores['quant'] * 30 + scores['diversity'] * 15 + scores['word_count'] * 15 + scores['skills_cat'] * 10)
    return {
        "score": round(final_score, 2), 
        "feedback": feedback,
        "metrics": {
            "keyword": round(scores['keyword'] * 100, 2),
            "quant": round(scores['quant'] * 100, 2),
            "diversity": round(scores['diversity'] * 100, 2),
            "word_count": round(scores['word_count'] * 100, 2),
            "skills": round(scores['skills_cat'] * 100, 2)
        }
    }

def analyze_jd(jd_text: str) -> Optional[Dict[str, Any]]:
    """Leverages LLM intelligence to deep-scan a Job Description and extract structural requirements.

    This function acts as the "Discovery Agent" in the pipeline. It parses the JD 
    to find not just keywords, but the company's implicit priorities and desired 
    experience profile.

    Features:
    - Model Fallback: Tries multiple Gemini versions to ensure uptime and quality.
    - JSON Extraction: Forces the LLM to output structured data for programmatic use.

    Args:
        jd_text (str): The raw text of the job description provided by the user.

    Returns:
        Optional[Dict[str, Any]]: A dictionary containing 'must_have_skills', 
            'nice_to_have_skills', 'ideal_job_title', etc. Returns None on total failure.
    """
    prompt = f"""
    You are an expert Talent Acquisition Specialist and ATS Algorithm Expert.
    Analyze the following Job Description and extract structural data for a resume builder.
    
    JOB DESCRIPTION:
    {jd_text}
    
    OUTPUT FORMAT (JSON):
    {{
        "ideal_job_title": "...",
        "must_have_skills": ["skill 1", "skill 2", ...],
        "nice_to_have_skills": ["...", ...],
        "experience_requirements": "...",
        "industry_keywords": ["keyword 1", "keyword 2", ...],
        "top_3_priorities": ["What the company cares about most"]
    }}
    
    Focus on extracting hard technical skills, specific tools, and strategic keywords that an ATS would scan for.
    """
    for model_name in MODELS_TO_TRY:
        try:
            model = genai.GenerativeModel(model_name, generation_config={"response_mime_type": "application/json"})
            response = model.generate_content(prompt)
            if not response.text:
                print(f">>> EMPTY RESPONSE from {model_name}")
                continue
            return json.loads(response.text)
        except Exception as e:
            print(f">>> Gemini Error with {model_name}: {e}")
            continue
    return None


def tailor_resume_attempt(master_data: Dict[str, Any], job_description: str, analysis: Dict[str, Any], feedback: str = None) -> Optional[Dict[str, Any]]:
    """Executes a single engineering attempt to transform a master resume into a tailored version.

    This is the core "Synthesis Agent". It uses "Governance Prompting" — a strict 
    set of rules that prevent hallucinations while maximizing keyword absorption.

    Governance Rules Applied:
    - Word count enforcement (20-25 words per bullet).
    - Hard number requirement for every bullet point.
    - Zero action verb repetition across all entries.
    - Fresh/Entry-level tone specifically optimized for new grads.

    Args:
        master_data (Dict[str, Any]): The candidate's comprehensive master resume.
        job_description (str): the target JD text.
        analysis (Dict[str, Any]): The discovering requirements from analyze_jd.
        feedback (str, optional): Corrective feedback from the ATS scorer if this 
            is a retry attempt. This allows the model to "learn" from its previous score.

    Returns:
        Optional[Dict[str, Any]]: The tailored resume JSON object, or None if the 
            LLM fails to produce valid JSON.
    """
    target_title = analysis.get("ideal_job_title", "Target Role")
    keywords = analysis.get("must_have_skills", [])
    industry_kw = analysis.get("industry_keywords", [])
    top_3 = analysis.get("top_3_priorities", [])

    prompt = f"""
You are an elite, ATS-optimizing Resume Writer specializing in FRESHER / ENTRY-LEVEL candidates.


           TWO-TIER HONESTY RULE — READ CAREFULLY             

  TIER 1 — STRICT (Data Integrity):                           
    • USE USER DATA AS-IS: Your primary task is to maintain   
      the user's original company names, roles, and metrics.  
    • DO NOT rewrite the meaning of bullets. Simply rephrase  
      them to be exactly 20-25 words and ensure each starts   
      with a unique action verb.                              
    • QUANTIFICATION: Use the user's provided metrics. If the 
      user provided a number, DO NOT change it.               
                                                              
  TIER 2 — FLEXIBLE (Skills, Summary, Core Competencies):     
    • The Skills section, Summary, and Core Competencies ARE  
      the keyword absorption layer. You MUST include every    
      must-have and nice-to-have JD keyword here — even if    
      the candidate only has academic/theoretical exposure.   
    • A skill appearing in the Skills section signals         
      awareness and readiness — this is standard practice     
      for fresher resumes and is ATS-expected.                
══════════════════════════════════════════════════════════════
TARGET ROLE   : {target_title}
MUST-HAVE KW  : {', '.join(keywords)}
INDUSTRY KW   : {', '.join(industry_kw)}
TOP PRIORITIES: {'; '.join(top_3)}
══════════════════════════════════════════════════════════════

━━━━━━━━━━━━━━━━━━━  SECTION-BY-SECTION RULES  ━━━━━━━━━━━━━━━━━━━

[SUMMARY — 3 to 4 lines, approx 60-80 words]
  • This is a FRESHER resume. The summary must read like an ambitious student/new-grad,
    NOT a 5-year veteran.
  • MANDATORY — Sentence 1 MUST open with the exact target job title "{target_title}" and 
    the candidate's degree context. (Do NOT include name). Example: "Aspiring {target_title} with a strong 
    foundation in..." — The title "{target_title}" MUST appear verbatim in Sentence 1.
  • Sentence 2: What specific value they bring to THIS role (tie to top_3_priorities).
  • Content: Naturally weave in 2–3 of the most critical JD must-have keywords 
    (from MUST-HAVE KW list above) into the paragraph — do NOT list them, blend them.
  • Highlight a high-impact project outcome and specific value aligned with top_3_priorities.
  • Format: Write a cohesive paragraph that is 3 to 4 lines long (about 70 words).
  • Hard limit: 80 words total. If it exceeds 80 words, cut ruthlessly.
  • FORBIDDEN in summary: lists, semicolons, bullet-style phrases, the word "proficient",
    "passionate", "detail-oriented", "team player", "self-motivated", "hardworking".
  • Do NOT mention every project or every skill here — that is what the other sections are for.

[TECHNICAL SKILLS — exactly 4 to 5 domain categories]  ← PRIMARY KEYWORD ABSORPTION LAYER
  • MANDATORY: One of these categories MUST be "Soft Skills" (or similar).
  • Category names must be concrete (e.g., "ML & AI Frameworks", "Data Engineering",
    "Soft Skills", "Languages & Databases"). No generic names like "Others".
  • Each category: 4 to 7 skills, comma-separated.
  • MANDATORY COVERAGE: Every single high-level keyword and soft skill from MUST-HAVE KW 
    and INDUSTRY KW above MUST appear in at least one skills category. Do not skip any.
  • Skills from master data take priority placement; JD-only keywords fill remaining slots.
  • Do NOT repeat a skill across categories.

[EXPERIENCE — exactly 2 entries, most relevant to JD]
  • Pick the 2 internships/roles closest to the target title.
  • ORDER: Always output the MOST RECENT experience FIRST (descending chronological order).
  • Each entry: exactly 3 bullet points.
  • See UNIVERSAL BULLET RULES below.

[PROJECTS — exactly 3 Projects, exactly 4 bullet points each]
  • Pick the 3 projects whose tech_stack and narrative best match the JD keywords.
  • Each project header: list exactly 5 technologies (most ATS-relevant ones).
  • Each entry: EXACTLY 4 bullet points.
  • ONE-LINE CONSTRAINT: Every bullet must be exactly ONE LINE long (approx 18-22 words).
  • BULLET SYNTHESIS: You MUST build each bullet by combining one 'core_action' with 
    one 'quantified_impact' from the project data. Do not just copy-paste; rephrase 
    to meet the strict 18-22 word requirement for single-line display.
  • See UNIVERSAL BULLET RULES below.

[EDUCATION]
  • Copy verbatim from master data. Do not add or remove fields.

[PERSONAL INFO]
  • Copy ALL fields verbatim from master data: name, phone, email, linkedin, github, portfolio, kaggle.
  • Never omit kaggle even if not mentioned in the JD.

[CERTIFICATIONS]
  • Copy verbatim from master data as a flat list.

━━━━━━━━━━━━━━━━━━━  UNIVERSAL BULLET RULES (NON-NEGOTIABLE)  ━━━━━━━━━━━━━━━━━━━

RULE B-1  WORD COUNT: Every bullet must be between 18 and 22 words (inclusive).
          This ensures every bullet stays on exactly ONE LINE in the PDF. Bullets outside 
          this range will fail the ATS scorer and break the page layout.

RULE B-2  QUANTIFICATION: Every bullet must contain at least one hard number, percentage,
          dollar figure, time-saving metric, or dataset scale (e.g., 500K rows, 30%, 4×).
          "Improved performance" is INVALID. "Improved inference speed by 38%" is VALID.

RULE B-3  ZERO VERB REPETITION: Every single bullet point (all 18 across Experience & 
          Projects) MUST start with a UNIQUE action verb. You are FORBIDDEN from using 
          the same verb twice.
          
          • BANNED REPEATS: Never reuse 'Built', 'Developed', 'Integrated', 'Designed', 
            'Implemented', or 'Created'. Use them once max, or find better synonyms.
            
          • SUGGESTED UNIQUE VERBS (Use these to ensure diversity):
            - Technical/Architecture: Architected, Orchestrated, Engineered, Deployed, 
              Refactored, Parallelized, Migrated, Calibrated, Benchmarked, Vectorized.
            - Analysis/Math: Synthesized, Validated, Audited, Quantified, Extracted, 
              Extrapolated, Modeled, Forecasted, Standardized, Reconstructed.
            - Performance: Optimized, Streamlined, Scaled, Accelerated, Reduced, 
              Maximized, Boosted, Diminished, Mitigated, Expedited.
            - Workflow: Automated, Centralized, Codified, Synchronized, Translated, 
              Visualized, Mapped, Scripted, Documented, Initialized.
            - Leadership: Catalyzed, Spearheaded, Championed, Authored, Directed.

RULE B-4  KEYWORD INJECTION: At least 1 must-have keyword from the JD must appear
          naturally inside each experience/project bullet. Do not force — rewrite the
          bullet context to make it fit naturally.

RULE B-5  NO VAGUE LANGUAGE: Ban these phrases entirely —
          "various", "several", "multiple tasks", "worked on", "helped with",
          "responsible for", "assisted in", "exposure to", "familiar with".

RULE B-6  GROWTH SIGNAL (at least 1 bullet per section): Show technical evolution —
          migration, version upgrades, scale-up, or before/after optimization.
          Example: "Migrated batch pipeline from pandas to Dask, cutting memory usage by 62%
          across 1.2M-row datasets in production."

RULE B-7  NO SUMMARY BLEED: Bullet points must describe concrete actions and results.
          They must NOT be a rewording of the summary or a generic statement of skills.

RULE B-8  METADATA PRESERVATION: You MUST include the "github_link" for every project 
          provided in the input data. Do NOT omit this field. If no link is provided, 
          use an empty string "".

━━━━━━━━━━━━━━━━━━━  SELF-VALIDATION CHECKLIST (run before outputting)  ━━━━━━━━━━━━━━━━━━━

Before writing the final JSON, mentally verify:
  [ ] Summary Sentence 1 contains the EXACT job title "{target_title}" — non-negotiable.
  [ ] Summary naturally includes 2-3 JD must-have keywords.
  [ ] Summary is 3-4 lines (approx 70 words) and follows the fresher tone.
  [ ] Exactly 6 core competencies — at least 4 are direct JD keywords.
  [ ] Exactly 4–5 skill categories.
  [ ] EVERY keyword from MUST-HAVE KW list appears in at least one skills category.
  [ ] Experience and project bullets contain ZERO fabricated tools, companies, or metrics.
  [ ] Exactly 2 experience entries, 3 bullets each, MOST RECENT FIRST.
  [ ] Exactly 3 project entries, 5 tech tags and 4 bullets each.
  [ ] Every bullet is exactly 18–22 words (to ensure single-line display).
  [ ] Every bullet has a quantified metric.
  [ ] No action verb is repeated across all 18 bullets.
  [ ] No fabricated data — everything traces back to master_data.

{f"⚠️  PREVIOUS ATTEMPT FAILED. FIX THESE ISSUES FIRST: {feedback}" if feedback else ""}

━━━━━━━━━━━━━━━━━━━  INPUT DATA  ━━━━━━━━━━━━━━━━━━━

MASTER RESUME JSON:
{json.dumps(master_data, indent=2)}

JOB DESCRIPTION:
{job_description}

━━━━━━━━━━━━━━━━━━━  OUTPUT FORMAT (strict JSON, no markdown)  ━━━━━━━━━━━━━━━━━━━

Return ONLY this JSON object with no extra text, no markdown fences:
{{
    "personal_info": {{
        "name": "...",
        "location": "City, State",
        "phone": "...",
        "email": "...",
        "linkedin": "...",
        "github": "...",
        "portfolio": "...",
        "kaggle": "..."
    }},
    "summary": "[MUST open with target job title: {target_title}] Cohesive 3-4 line paragraph, approx 70 words, fresher tone. MANDATORY: Absorb ALL soft skills and key metrics from JD here.",
    "skills": {{
        "Domain Category 1": ["Skill A", "Skill B", "Skill C"],
        "Domain Category 2": ["Skill D", "Skill E", "Skill F"],
        "Domain Category 3": ["Skill G", "Skill H", "Skill I"],
        "Domain Category 4": ["Skill G", "Skill H", "Skill I"]
    }},
    "experience": [
        {{
            "company": "Company Name",
            "role": "Exact Role Title",
            "start_date": "Mon YYYY",
            "end_date": "Mon YYYY or Present",
            "bullet_points": [
                "Action verb + task + quantified result, 20-25 words total.",
                "Action verb + task + quantified result, 20-25 words total.",
                "Action verb + task + quantified result, 20-25 words total."
            ]
        }}
    ],
    "projects": [
        {{
            "name": "Project Name",
            "github_link": "github.com/...",
            "technologies": ["Tech1", "Tech2", "Tech3", "Tech4", "Tech5"],
            "bullet_points": [
                "Action verb + task + quantified result, 18-22 words total (Exactly 1 line).",
                "Action verb + task + quantified result, 18-22 words total (Exactly 1 line).",
                "Action verb + task + quantified result, 18-22 words total (Exactly 1 line).",
                "Action verb + task + quantified result, 18-22 words total (Exactly 1 line)."
            ]
        }}
    ],
    "education": [
        {{
            "institution": "Full Institution Name",
            "degree": "Full Degree Name",
            "location": "City, State",
            "start_date": "Mon YYYY",
            "end_date": "Mon YYYY",
            "cgpa": "Value from 'gpa' or 'cgpa' field"
        }}
    ],
    "certifications": ["Cert Name — Issuer", "Cert Name — Issuer"]
}}
"""
    for model_name in MODELS_TO_TRY:
        try:
            model = genai.GenerativeModel(model_name, generation_config={"response_mime_type": "application/json"})
            return json.loads(model.generate_content(prompt).text)
        except: continue
    return None

@app.post("/api/tailor")
async def tailor_endpoint(master_json: str = Form(...), jd: str = Form(...)):
    """Primary API endpoint for the AI Tailoring Pipeline.

    This function coordinates the "Agentic Loop":
    1. JD Discovery: Analyzes the target job requirements.
    2. Synthesis: Attempts to generate an optimized resume.
    3. Evaluation: Scores the resume using the internal ATS Scorer.
    4. Iteration: If the score is below the 95% threshold, it provides feedback 
       to the Synthesis agent and tries again (up to 3 times).
    5. Finalization: Cleans up data and returns the highest-scoring version.

    Complexity: O(N) where N is max_attempts.
    Network: Multiple calls to Google Gemini API.

    Args:
        master_json (str): User's profile data in JSON format (received as Form data).
        jd (str): The raw text of the job description.

    Returns:
        dict: A success payload containing the best tailored JSON data, 
            the final ATS score, and detailed performance metrics.
    """
    async def event_generator():
        try:
            yield json.dumps({"step": "step-wait"}) + "\n"
            master_data = json.loads(master_json)
            yield json.dumps({"step": "step-load"}) + "\n"
            
            analysis = analyze_jd(jd)
            if not analysis: raise Exception("JD Analysis failed.")
            yield json.dumps({"step": "step-parse"}) + "\n"
                
            best_data = None
            max_attempts = 1
            current_attempt = 1
            feedback = None
            final_report = {"score": 0}

            while current_attempt <= max_attempts:
                yield json.dumps({"step": "step-core"}) + "\n"
                tailored_data = tailor_resume_attempt(master_data, jd, analysis, feedback)
                if not tailored_data:
                    current_attempt += 1
                    continue
                
                yield json.dumps({"step": "step-exp"}) + "\n"
                report = score_resume_internal(tailored_data, analysis)
                final_report = report
                yield json.dumps({"step": "step-other"}) + "\n"
            
                # HARD-CODED FAIL-SAFE: Enforce exact layout count
                tailored_data['experience'] = tailored_data.get('experience', [])[:2]
                tailored_data['projects'] = tailored_data.get('projects', [])[:3]
                
                for exp in tailored_data.get('experience', []):
                    exp['bullet_points'] = exp['bullet_points'][:3]
                for proj in tailored_data.get('projects', []):
                    proj['bullet_points'] = proj['bullet_points'][:4]

                if report['score'] >= 95.0:
                    best_data = tailored_data
                    break
                else:
                    feedback = ". ".join(report['feedback'])
                    if not best_data or report['score'] > final_report.get('score', 0):
                        best_data = tailored_data
                    current_attempt += 1

            if not best_data: raise Exception("Iterative tailoring failed.")

            yield json.dumps({
                "success": True,
                "data": best_data,
                "score": final_report['score'],
                "metrics": final_report.get('metrics', {})
            }) + "\n"
        except Exception as e:
            yield json.dumps({"success": False, "error": str(e)}) + "\n"

    return StreamingResponse(event_generator(), media_type="application/x-ndjson")

@app.post("/api/download")
async def download_pdf_direct(data: dict, template: str = "classic"):
    """API endpoint to generate and stream a PDF resume.

    Args:
        data (dict): The finalized resume JSON data.

    Returns:
        StreamingResponse: A PDF file stream with appropriate headers.
    """
    try:
        buffer = io.BytesIO()
        create_pdf(data, buffer, template)
        buffer.seek(0)
        
        filename = f"Tailored_Resume_{datetime.now().strftime('%Y%m%d%H%M')}.pdf"
        return StreamingResponse(
            buffer, 
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

def create_docx(data: dict, output_buffer):
    """Generates a DOCX resume from structured JSON data."""
    doc = Document()
    
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    p_info = data.get("personal_info", {})
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(p_info.get("name", "Name"))
    run.bold = True
    run.font.size = Pt(22)

    contact_parts = []
    if p_info.get("email"): contact_parts.append(p_info["email"])
    if p_info.get("phone"): contact_parts.append(p_info["phone"])
    if p_info.get("location"): contact_parts.append(p_info["location"])
    if p_info.get("linkedin"): contact_parts.append(p_info["linkedin"])
    if p_info.get("github"): contact_parts.append(p_info["github"])
    
    if contact_parts:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_p.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(10)

    def add_header(text):
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    if data.get("summary"):
        add_header("Professional Summary")
        p = doc.add_paragraph(data["summary"])
        p.style.font.size = Pt(10)

    if data.get("skills"):
        add_header("Technical Skills")
        skills = data["skills"]
        skills_p = doc.add_paragraph()
        skills_p.style.font.size = Pt(10)
        if isinstance(skills, dict):
            for k, v in skills.items():
                if v:
                    r = skills_p.add_run(f"{k.replace('_', ' ').title()}: ")
                    r.bold = True
                    r.font.size = Pt(10)
                    r2 = skills_p.add_run(", ".join(v) + "\n")
                    r2.font.size = Pt(10)
        elif isinstance(skills, list):
            r = skills_p.add_run(", ".join(skills))
            r.font.size = Pt(10)

    if data.get("experience"):
        add_header("Experience")
        for exp in data["experience"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r1 = p.add_run(exp.get("company", ""))
            r1.bold = True
            r1.font.size = Pt(11)
            
            r2 = p.add_run(f" | {exp.get('role', '')}")
            r2.italic = True
            r2.font.size = Pt(11)
            
            date_str = ""
            if exp.get("start_date") and exp.get("end_date"):
                date_str = f"{exp['start_date']} - {exp['end_date']}"
            if date_str:
                r3 = p.add_run(f"  [{date_str}]")
                r3.font.size = Pt(10)
            
            for bullet in exp.get("bullet_points", []):
                bp = doc.add_paragraph(bullet, style='List Bullet')
                bp.paragraph_format.space_after = Pt(2)
                for run in bp.runs:
                    run.font.size = Pt(10)

    if data.get("projects"):
        add_header("Projects")
        for proj in data["projects"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r1 = p.add_run(proj.get("name", ""))
            r1.bold = True
            r1.font.size = Pt(11)
            
            if proj.get("tech_stack"):
                r2 = p.add_run(f" | {', '.join(proj['tech_stack'])}")
                r2.italic = True
                r2.font.size = Pt(11)
            
            for bullet in proj.get("bullet_points", []):
                bp = doc.add_paragraph(bullet, style='List Bullet')
                bp.paragraph_format.space_after = Pt(2)
                for run in bp.runs:
                    run.font.size = Pt(10)

    if data.get("education"):
        add_header("Education")
        for edu in data["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r1 = p.add_run(edu.get("institution", ""))
            r1.bold = True
            r1.font.size = Pt(11)
            
            r2 = p.add_run(f" | {edu.get('degree', '')}")
            r2.font.size = Pt(11)
            
            if edu.get("end_date"):
                r3 = p.add_run(f" ({edu.get('end_date')})")
                r3.font.size = Pt(11)

    doc.save(output_buffer)

@app.post("/api/download_docx")
async def download_docx_direct(data: dict):
    """API endpoint to generate and stream a DOCX resume."""
    try:
        buffer = io.BytesIO()
        create_docx(data, buffer)
        buffer.seek(0)
        
        filename = f"Tailored_Resume_{datetime.now().strftime('%Y%m%d%H%M')}.docx"
        return StreamingResponse(
            buffer, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/sample_master")
async def get_sample():
    """Returns the standardized sample_master.json for user reference."""
    try:
        sample_path = os.path.join(STATIC_DIR, "sample_master.json")
        with open(sample_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        return {"error": f"Could not load sample: {str(e)}"}

@app.get("/")
async def home_page():
    return FileResponse(os.path.join(STATIC_DIR, "index.html"))

@app.get("/input")
async def input_page():
    return FileResponse(os.path.join(STATIC_DIR, "input.html"))

@app.get("/result")
async def result_page():
    return FileResponse(os.path.join(STATIC_DIR, "result.html"))

@app.get("/guide")
async def guide_page():
    return FileResponse(os.path.join(STATIC_DIR, "guide.html"))

@app.get("/privacy")
async def privacy_page():
    return FileResponse(os.path.join(STATIC_DIR, "privacy.html"))

@app.get("/about")
async def about_page():
    return FileResponse(os.path.join(STATIC_DIR, "about.html"))

@app.get("/editor")
async def editor_page():
    return FileResponse(os.path.join(STATIC_DIR, "editor.html"))

@app.get("/robots.txt")
async def robots_txt():
    return FileResponse(os.path.join(STATIC_DIR, "robots.txt"))

@app.get("/sitemap.xml")
async def sitemap_xml():
    return FileResponse(os.path.join(STATIC_DIR, "sitemap.xml"))

@app.get("/seo.json")
async def seo_json():
    return FileResponse(os.path.join(STATIC_DIR, "seo.json"))

@app.get("/manifest.json")
async def manifest_json():
    return FileResponse(os.path.join(STATIC_DIR, "manifest.json"))

@app.get("/browserconfig.xml")
async def browserconfig_xml():
    return FileResponse(os.path.join(STATIC_DIR, "browserconfig.xml"))

@app.get("/api/seo/sitemap")
async def get_seo_sitemap():
    return {
        "success": True,
        "base_url": "https://tailored-resume-ai.vercel.app",
        "seo_config": "https://tailored-resume-ai.vercel.app/seo.json",
        "sitemap": "https://tailored-resume-ai.vercel.app/sitemap.xml",
        "robots": "https://tailored-resume-ai.vercel.app/robots.txt",
        "author": "Nikhil Pitroda",
        "urls": [
            {"path": "/", "name": "WorkSpace (Home)", "priority": 1.0, "changefreq": "monthly"},
            {"path": "/guide", "name": "Master Resume Guide", "priority": 0.8, "changefreq": "monthly"},
            {"path": "/about", "name": "Project Architecture", "priority": 0.7, "changefreq": "monthly"},
            {"path": "/input", "name": "AI Tailoring Dashboard", "priority": 0.6, "changefreq": "monthly"},
            {"path": "/privacy", "name": "Security Policy", "priority": 0.3, "changefreq": "yearly"}
        ]
    }

# Mount static files AFTER routes
app.mount("/", StaticFiles(directory=STATIC_DIR), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
